import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_template():
    doc = Document()
    
    # Estilos básicos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(14)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Endereçamento
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DE DIREITO DO JUIZADO ESPECIAL CÍVEL")
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph("\n")
    doc.add_paragraph("\n")
    
    # Qualificação
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("Processo nº: {{NUMERO_PROCESSO}}\n\n")
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("{{NOME_CLIENTE}}").bold = True
    p.add_run(", já qualificado(a) nos autos do processo em epígrafe, vem, respeitosamente, à presença de Vossa Excelência, por seus advogados que esta subscrevem, apresentar sua ")
    p.add_run("{{TIPO_PECA}}").bold = True
    p.add_run(", pelos motivos de fato e de direito a seguir expostos:\n")
    
    # Fatos
    p = doc.add_paragraph("I. DOS FATOS\n")
    p.runs[0].bold = True
    p = doc.add_paragraph("{{RESUMO_FATOS}}\n")
    
    # Direito
    p = doc.add_paragraph("II. DO DIREITO\n")
    p.runs[0].bold = True
    p = doc.add_paragraph("{{FUNDAMENTACAO_JURIDICA}}\n")
    
    # Pedidos
    p = doc.add_paragraph("III. DOS PEDIDOS\n")
    p.runs[0].bold = True
    p = doc.add_paragraph("Ante o exposto, requer a Vossa Excelência:\n")
    p = doc.add_paragraph("{{PEDIDOS}}\n")
    
    p = doc.add_paragraph("Termos em que, pede deferimento.\n")
    p = doc.add_paragraph("Local, Data.\n")
    p = doc.add_paragraph("_________________________________\nAdvogado(a)\nOAB/UF nº XXXXX")
    
    os.makedirs('templates', exist_ok=True)
    doc.save('templates/modelo_defesa.docx')
    print("Template criado com sucesso em templates/modelo_defesa.docx")

if __name__ == "__main__":
    create_template()
